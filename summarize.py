import json
import os
from pydub import AudioSegment
from docx import Document
from transformers import pipeline
import whisper
from llama_cpp import Llama

# Initialize models
llm = Llama(model_path="OpenBioLLM-Llama3-8B.Q4_K_M.gguf", chat_format="llama-2", cpu_threads=4, n_ctx=4096)
transcriber_model = whisper.load_model("base")
summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
classifier = pipeline("zero-shot-classification", model="facebook/bart-large-mnli")

# Labels for medical text categorization
labels = ["Vitals", "Pain", "Skin", "Respiratory", "Neuro", "Musculoskeletal", "Gastrointestinal", "Medication", "Summary"]

# Convert non-wav audio formats to wav
def convert_to_wav(audio_path):
    file_extension = audio_path.split(".")[-1].lower()
    if file_extension != "wav":
        audio = AudioSegment.from_file(audio_path)
        wav_path = "temp_audio.wav"
        audio.export(wav_path, format="wav")
        return wav_path
    return audio_path

# Transcribe audio and summarize the transcription
def transcribe_and_summarize(audio_path):
    audio_path = convert_to_wav(audio_path)
    transcription_result = transcriber_model.transcribe(audio_path)
    transcription_text = transcription_result["text"]

    # Summarize transcription
    summarized_text = summarizer(transcription_text, max_length=300, min_length=100, do_sample=False)[0]["summary_text"]

    if audio_path == "temp_audio.wav":
        os.remove(audio_path)

    return transcription_text, summarized_text

# Structure summarized text into labeled categories
def structure_text(summarized_text):
    segments = summarized_text.split(". ")
    structured_text = {label: [] for label in labels}

    for segment in segments:
        if segment.strip():
            classification = classifier(segment, candidate_labels=labels)
            label = classification["labels"][0]
            structured_text[label].append(segment)

    return structured_text

# Create a formatted table for vitals in the document
def add_vitals_table(doc, vitals_data):
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Time'
    hdr_cells[1].text = 'Blood Pressure'
    hdr_cells[2].text = 'Temperature'
    hdr_cells[3].text = 'Respiration'
    hdr_cells[4].text = 'Pulse'

    for row_data in vitals_data:
        row_cells = table.add_row().cells
        for i, item in enumerate(row_data):
            row_cells[i].text = str(item)

# Create a nursing note document
def create_nursing_note_doc(transcription, summary, structured_text, output_path="Nursing_Note.docx"):
    doc = Document()
    doc.add_heading("Nursing Note", 0)

    # Add Transcription section
    doc.add_heading("Transcription:", level=1)
    doc.add_paragraph(transcription)

    # Add Summary section
    doc.add_heading("Summary:", level=1)
    doc.add_paragraph(summary)

    # Process each category and add it to the document
    for section, content in structured_text.items():
        if content:
            doc.add_heading(section, level=1)

            # For vitals, add to table
            if section == "Vitals":
                vitals_data = []
                for line in content:
                    if any(x in line for x in ["AM", "PM"]):
                        parts = line.split(",")
                        if len(parts) >= 5:
                            time, bp, temp, resp, pulse = parts[:5]
                            vitals_data.append((time.strip(), bp.strip(), temp.strip(), resp.strip(), pulse.strip()))
                add_vitals_table(doc, vitals_data)
            else:
                for line in content:
                    # Checkbox format for specific sections
                    if section in ["Pain", "Skin", "Respiratory"]:
                        if "No" in line:
                            doc.add_paragraph("☐ " + line)
                        else:
                            doc.add_paragraph("☑ " + line)
                    else:
                        doc.add_paragraph(line)

    # Save the document
    doc.save(output_path)

# Main function to integrate all parts
def main(audio_path):
    transcription_text, summarized_text = transcribe_and_summarize(audio_path)
    structured_text = structure_text(summarized_text)
    create_nursing_note_doc(transcription_text, summarized_text, structured_text)
    print("Nursing note created successfully.")

# Input audio path
audio_path = "/content/Patient_Report_Audio_2.wav"
main(audio_path)
